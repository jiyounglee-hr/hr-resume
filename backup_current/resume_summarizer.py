import os
import pandas as pd
from pathlib import Path
import openai
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path
import tempfile
from openai import OpenAI
import asyncio
import aiohttp
from functools import lru_cache
from docx import Document
from docx.shared import Inches

# 환경 변수 로드
load_dotenv()

class ResumeSummarizer:
    def __init__(self):
        self.client = OpenAI()
        self.model = "gpt-3.5-turbo"  # 또는 "gpt-4-turbo-preview"
        
        self.jd_analysis_prompt = """
        다음 이력서와 JD를 분석하여 핵심적인 적합성 분석을 제공해주세요.
        이름은 한 번만 언급하고, 이후에는 '지원자'로 표현해주세요:
        
        이력서:
        {resume}
        
        JD:
        {jd}
        
        다음 항목들을 중심으로 분석해주세요:
        1. 필수자격 요건 충족도
        2. 우대사항 충족도
        3. 경력 적합성
        """
        
        self.core_values_prompt = """
        다음 이력서를 분석하여 각 핵심가치별로 평가해주세요.
        이름은 언급하지 말고 '지원자'로 표현해주세요:
        
        핵심가치:
        1. 도전정신 (두려워 말고 시도합니다)
        2. 책임감 (대충은 없습니다)
        3. 협력 (동료와 협업합니다)
        4. 전문성 (능동적으로 일합니다)
        
        이력서:
        {resume}
        """
        
        self.interview_prompt = """
        다음 이력서와 JD를 바탕으로 면접 질문을 생성해주세요.
        이름을 언급하지 말고 '지원자'로 표현해주세요.

        1. 직무 관련 질문 (5개):
        - 지원자의 경력과 JD 요구사항을 연결하여 구체적인 질문 생성
        - 실제 업무 수행 능력을 파악할 수 있는 상황 기반 질문 포함
        - 전문성과 경험을 검증할 수 있는 기술적 질문 포함

        2. 핵심가치 관련 질문 (4개):
        - 도전정신: 어려움을 극복한 경험 관련 질문
        - 책임감: 맡은 일을 완수한 경험 관련 질문
        - 협력: 팀워크와 협업 경험 관련 질문
        - 전문성: 자기주도적 업무 수행 경험 관련 질문

        이력서:
        {resume}
        
        JD:
        {jd}

        답변 형식:
        [직무 관련 질문]
        1. 질문1
        2. 질문2
        3. 질문3
        4. 질문4
        5. 질문5

        [핵심가치 관련 질문]
        1. [도전정신] 질문1
        2. [책임감] 질문2
        3. [협력] 질문3
        4. [전문성] 질문4
        """

    def extract_text_from_pdf_with_ocr(self, pdf_path):
        """PDF에서 OCR을 사용하여 텍스트를 추출합니다."""
        try:
            # 먼저 일반적인 방법으로 텍스트 추출 시도
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text

            # 추출된 텍스트가 없으면 OCR 시도
            if not text.strip():
                print("일반 텍스트 추출 실패, OCR 시작...")
                # PDF를 이미지로 변환
                images = convert_from_path(pdf_path)
                text = ""
                for image in images:
                    # 이미지에서 텍스트 추출 (한국어 포함)
                    text += pytesseract.image_to_string(image, lang='kor+eng') + "\n"

            return text.strip()
        except Exception as e:
            print(f"OCR 처리 중 오류 발생: {str(e)}")
            return None

    def read_resume(self, resume_path):
        """이력서 파일을 읽어옵니다."""
        try:
            if resume_path.endswith('.pdf'):
                try:
                    reader = PdfReader(resume_path)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    if not text.strip():
                        return "PDF에서 텍스트를 추출할 수 없습니다. 스캔된 PDF일 수 있습니다."
                    return text
                except Exception as pdf_error:
                    print(f"PDF 처리 중 오류 발생: {str(pdf_error)}")
                    return None
            elif resume_path.endswith('.docx'):
                return "Word 문서 처리는 아직 지원되지 않습니다."
            else:
                # 다양한 인코딩 시도
                encodings = ['utf-8', 'euc-kr', 'cp949']
                for encoding in encodings:
                    try:
                        with open(resume_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                return "파일 인코딩을 확인할 수 없습니다."
        except Exception as e:
            print(f"이력서 파일 읽기 오류: {str(e)}")
            return None

    def analyze_jd_fit(self, resume_text, jd_text):
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant analyzing resumes in Korean."},
                    {"role": "user", "content": f"""
이력서와 JD를 비교 분석하여 다음 형식으로 작성해주세요:

<h2>1. 직무상 강점</h2>
<ul>
[강점 리스트 작성]
</ul>

<h2>2. 직무상 약점 (이력서에 서술되지 않은 경력)</h2>
<ul>
[약점 리스트 작성]
</ul>

이력서:
{resume_text}

JD:
{jd_text}
"""}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Error in analyze_jd_fit: {str(e)}")
            return "분석 중 오류가 발생했습니다."

    def extract_name(self, resume_text):
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Extract only the Korean name from the resume text, without any additional text."},
                    {"role": "user", "content": f"""
다음 이력서에서 지원자의 한국어 이름만 추출해주세요:
- 설명이나 부가 텍스트 없이 이름만 반환
- 예시: "이건훈"

이력서:
{resume_text}
                """}
                ]
            )
            name = response.choices[0].message.content.strip()
            # 불필요한 텍스트 제거
            name = name.replace('지원자의 한국어 이름은 ', '').replace('"', '').replace('입니다', '')
            name = name.replace('[', '').replace(']', '').strip()
            return name
        except Exception as e:
            return "지원자"

    def analyze_core_values(self, resume_text):
        """핵심가치 기반 분석을 수행합니다."""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": f"다음 이력서를 분석해주세요:\n{resume_text}"}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"분석 중 오류가 발생했습니다: {str(e)}"

    def generate_interview_questions(self, resume_text, jd_text):
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant generating interview questions in Korean."},
                    {"role": "user", "content": f"""
다음 이력서와 JD를 바탕으로 면접 질문을 생성해주세요:

이력서:
{resume_text}

JD:
{jd_text}

다음 형식으로 작성해주세요:

<div class="interview-guide">
    <h1>면접 질문 가이드</h1>

    <h2>💡 JD 기반 질문:</h2>    
    <ul>
     [이력서를 바탕으로 JD의 역량 및 문제해결 능력 검증을 위한 7-8개의 질문 생성]
    </ul>
    <h2>🚀 핵심가치 기반 질문:</h2>
    1.두려워 말고 시도합니다. (도전정신)
    <ul>
      [2-3개의 관련 질문을 생성]
    </ul>
    2.대충은 없습니다. (책임감)
    <ul>
      [2-3개의 관련 질문을 생성]
    </ul>
    3.동료와 협업합니다. (협력)
    <ul>
      [2-3개의 관련 질문을 생성]
    </ul>
    4.능동적으로 일합니다. (전문성)
    <ul>
      [2-3개의 관련 질문을 생성]
    </ul>
</div>
                """}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"면접 질문 생성 중 오류가 발생했습니다: {str(e)}"

    @lru_cache(maxsize=100)
    def generate_summary(self, resume_text, jd_text):
        try:
            name = self.extract_name(resume_text)
            title = f"[{name}]님의 이력서 분석 결과"
            jd_analysis = self.analyze_jd_fit(resume_text, jd_text)
            
            return f"<h1>{title}</h1>\n\n{jd_analysis}"
        except Exception as e:
            return f"분석 중 오류가 발생했습니다: {str(e)}"

    async def analyze_all(self, resume_text, jd_text):
        tasks = [
            self.analyze_jd_fit(resume_text, jd_text),
            self.analyze_core_values(resume_text),
            self.generate_interview_questions(resume_text, jd_text)
        ]
        results = await asyncio.gather(*tasks)
        return "\n\n".join(results)

def format_interview_questions(document):
    # 면접 질문 가이드 섹션
    document.add_heading('면접 질문 가이드', level=1)
    
    # JD 기반 질문
    document.add_paragraph('💡 JD 기반 질문:', style='Heading 2')
    jd_questions = [
        "의료기기 인허가 업무 경험에 대해 구체적으로 언급한 주요 성과 중 하나를 설명해주실 수 있나요?",
        "국가별 의료기기 규제 요구사항 분석 시 복잡한 상황에서 어떻게 해결해나갔는지에 대한 경험을 공유해주실 수 있을까요?",
        "품질시스템 심사 대응 과정에서 어려움을 겪었던 경험에 대해 이야기해주실 수 있을까요? 어떻게 대응했나요?",
        "의료기기 제품 인허가를 위한 교육 및 자격증 경험이 현재 업무에 어떻게 도움을 주고 있는지 설명해주세요.",
        "의료기기법/규격/가이드 변경사항 모니터링을 얼마나 효율적으로 수행할 수 있는지에 대해 이야기해주실 수 있을까요?"
    ]
    for q in jd_questions:
        p = document.add_paragraph()
        p.add_run(q)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 핵심가치 기반 질문
    document.add_paragraph('🚀 핵심가치 기반 질문:', style='Heading 2')
    
    # 1. 도전정신
    p = document.add_paragraph('1. 두려워 말고 시도합니다. (도전정신)')
    p.paragraph_format.left_indent = Inches(0.25)
    challenge_questions = [
        "과거 업무 중에 새로운 도전을 감행했던 경험에 대해 소개해주세요.",
        "새로운 프로젝트나 작업에 대해 도전적 제안을 한 적이 있는지 이야기해주실 수 있나요?"
    ]
    for q in challenge_questions:
        p = document.add_paragraph()
        p.add_run(q)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 2. 책임감
    p = document.add_paragraph('2. 대충은 없습니다. (책임감)')
    p.paragraph_format.left_indent = Inches(0.25)
    responsibility_questions = [
        "책임을 맡은 업무 중 특히 중요하다고 생각했던 일에 대해 자세히 설명해주세요.",
        "실수를 저질렀을 때 어떻게 책임을 지고 문제를 해결했던 경험이 있으신가요?"
    ]
    for q in responsibility_questions:
        p = document.add_paragraph()
        p.add_run(q)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 3. 협력
    p = document.add_paragraph('3. 동료와 협업합니다. (협력)')
    p.paragraph_format.left_indent = Inches(0.25)
    cooperation_questions = [
        "팀 내 외부팀과 협력하여 해결한 문제나 프로젝트가 있을까요? 해당 경험에 대해 이야기해주세요.",
        "자신이 협업을 통해 얻는 가장 큰 가치는 무엇이라고 생각하시나요?"
    ]
    for q in cooperation_questions:
        p = document.add_paragraph()
        p.add_run(q)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 4. 전문성
    p = document.add_paragraph('4. 능동적으로 일합니다. (전문성)')
    p.paragraph_format.left_indent = Inches(0.25)
    expertise_questions = [
        "자발적으로 전문성을 향상시키기 위해 노력한 경험을 공유해주세요.",
        "기술적 역량을 높이기 위해 어떤 노력을 하고 계신가요? 그 경험에 대해 이야기해주세요."
    ]
    for q in expertise_questions:
        p = document.add_paragraph()
        p.add_run(q)
        p.paragraph_format.left_indent = Inches(0.5)

def main():
    summarizer = ResumeSummarizer()
    
    # 예시 사용
    resume_path = "이력서"  # 실제 이력서 파일 경로로 수정 필요
    jd_text = "여기에 JD 텍스트를 입력하세요"
    
    summary = summarizer.generate_summary(resume_path, jd_text)
    print(summary)

if __name__ == "__main__":
    main() 