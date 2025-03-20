from flask import Flask, request, render_template, send_file, jsonify
from resume_summarizer import ResumeSummarizer
import os
from dotenv import load_dotenv
import socket
import PyPDF2
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import io
from flask_cors import CORS
import tempfile

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "https://[your-github-username].github.io"}})
summarizer = ResumeSummarizer()

def get_ip():
    # 현재 컴퓨터의 IP 주소 가져오기
    hostname = socket.gethostname()
    ip_address = socket.gethostbyname(hostname)
    return ip_address

def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': '이력서 파일이 필요합니다'}), 400
        
        resume_file = request.files['resume']
        jd_text = request.form.get('jd', '')
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            resume_file.save(temp_file.name)
            resume_text = summarizer.read_resume(temp_file.name)
        
        # 임시 파일 삭제
        os.unlink(temp_file.name)
        
        if resume_text is None:
            return jsonify({'error': '이력서 읽기 실패'}), 400
            
        # 분석 수행
        result = summarizer.generate_summary(resume_text, jd_text)
        
        return jsonify({'result': result})
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download-word', methods=['POST'])
def download_word():
    try:
        content = request.json.get('content')
        
        # Word 문서 생성
        doc = Document()
        
        # 가로 방향으로 설정
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)  # A4 가로
        section.page_height = Inches(8.27)  # A4 세로
        
        # 여백 설정 (좁게)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # 내용 추가 (공백 제거)
        sections = content.strip().split('\n\n')
        
        for section in sections:
            if section.strip():
                # 이모지가 있는 섹션 제목 처리
                if '📃' in section or '🚀' in section or '🎯' in section:
                    # 섹션 제목과 내용 분리
                    title_end = section.find(':')
                    if title_end != -1:
                        section_title = section[:title_end+1]
                        section_content = section[title_end+1:]
                        
                        # 섹션 제목 추가
                        heading = doc.add_paragraph()
                        heading.add_run(section_title).bold = True
                        
                        # 섹션 내용 추가 (앞뒤 공백 제거)
                        if section_content.strip():
                            doc.add_paragraph(section_content.strip())
                    else:
                        doc.add_paragraph(section.strip())
                else:
                    doc.add_paragraph(section.strip())
        
        # 스타일 적용
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
        
        # 메모리에 Word 파일 저장
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='이력서_분석_및_질문TIP.docx'
        )
        
    except Exception as e:
        return f"Word 파일 생성 중 오류가 발생했습니다: {str(e)}", 500

@app.route('/generate_questions', methods=['POST'])
def generate_questions():
    try:
        # 파일과 JD 텍스트 확인
        if 'file' not in request.files:
            return jsonify({'error': '이력서 파일(PDF)을 업로드해주세요.'}), 400
        
        file = request.files['file']
        jd_text = request.form.get('jd_text', '').strip()
        
        # 이력서 텍스트 추출
        resume_text = extract_text_from_pdf(file)
        
        # 면접 질문 생성
        questions = summarizer.generate_interview_questions(resume_text, jd_text)
        
        return jsonify({'html': questions})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    ip_address = get_ip()
    print(f"\n=== 서버 접속 정보 ===")
    print(f"내부 네트워크 접속 주소: http://{ip_address}:5000")
    print(f"로컬 접속 주소: http://localhost:5000")
    print("위의 주소로 접속하실 수 있습니다.")
    print("다른 사람들은 '내부 네트워크 접속 주소'로 접속하면 됩니다.")
    print("===================\n")
    
    # 모든 네트워크 인터페이스에서 접속 허용
    app.run(host='0.0.0.0', port=5000, debug=True) 